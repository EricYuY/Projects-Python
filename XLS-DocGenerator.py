###########################################################
############### GENERACION DE MEMORANDUMS #################
###########################################################

###########################################################
############### DESARROLLADO POR: INGENIERÍA KRCP #########
############### AUTOR: ERIC YU ############################
###########################################################

from docx import Document
import openpyxl
from pathlib import Path
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_LINE_SPACING
from docx.enum.text import WD_BREAK
from docx.shared import Pt

import tkinter as tk
from tkinter import ttk
from tkinter import messagebox
import time

# Con orden de compra
OS_arr=[]
Cliente_arr=[]
OC_arr=[]
Asunto_arr=[]

# Sin orden de compra
OS_arr_sinOC=[]
Cliente_arr_sinOC=[]
OC_arr_sinOC=[]
Asunto_arr_sinOC=[]

# Variables de interfaz (inputs)
fields = 'ARCHIVO (Ejm: JULIO.xlsx):', 'PERIODO (Ejm: JULIO 2021):', 'FILA INICIO:', 'FILA FIN:', "FECHA DE CIERRE (Ejm: 25/07/2021):"
field = []
text  = []
flag_finish = False

ruta=0
periodo=0
inicio=0
fin=0
fecha_cierre=0


def fetch(entries):
  field.clear()
  text.clear()
  for entry in entries:
    field.append(entry[0])
    if(entry[1].get()):
      text.append(entry[1].get())

def makeform(root, fields):
  entries = []
  for field in fields:
    row = tk.Frame(root)
    lab = tk.Label(row, width=30, text=field, anchor='w')
    ent = tk.Entry(row)
    row.pack(side=tk.TOP, fill=tk.X, padx=5, pady=5)
    lab.pack(side=tk.LEFT)
    ent.pack(side=tk.RIGHT, expand=tk.YES, fill=tk.X)
    entries.append((field, ent))
  return entries

def generar(root,entries):
  fetch(entries)
  if ((len(field)!=5)or(len(text)!=5)):
    messagebox.showinfo(title="Faltan datos", message="Ingrese datos")
    return;
  
  global ruta, periodo,inicio,fin,fecha_cierre
  ruta=text[0]
  periodo=text[1]
  inicio=text[2]
  fin=text[3]
  fecha_cierre=text[4]
  if((not inicio.isnumeric()) or (not fin.isnumeric())):
    messagebox.showinfo(title="Datos inválidos", message="Ingrese datos de fila inicio/fin correctos")
    return;
  barra_progreso(root)

def barra_progreso(root):
  root.destroy()
  root = tk.Tk()
  root.title("KRCP-MEMOS")
  root.geometry('300x100')
  progressbar = ttk.Progressbar(root,mode="indeterminate")
  progressbar.place(x=50, y=50, width=200)
  progressbar.start()

  generacion()
  
  b= tk.Button(root, text='Listo', command=(lambda e=progressbar: final_prog(e,root)))
  b.pack(side=tk.BOTTOM, padx=5, pady=5)
     
def final_prog(e,root):
  e.stop()
  root.destroy()
  root = tk.Tk()
  root.title("KRCP-MEMOS")
  root.geometry('300x60')

  value_label = ttk.Label(root, text="Proceso completado")
  value_label.pack(side=tk.TOP, padx=5, pady=5)
  b2 = tk.Button(root, text='Listo', command=(lambda e=root: cerrar(root)))
  b2.pack(side=tk.BOTTOM, padx=5, pady=5)

def cerrar(root):
  root.quit
  root.destroy()
  
def leer_xl(ruta,inicio,fin):
	wb_obj = openpyxl.load_workbook(ruta)
	#print(wb_obj)
	sheet=wb_obj.active

	for i in range(inicio,fin):
		OS=sheet["A{}".format(i)].value
		Cliente=sheet["B{}".format(i)].value
		OC=sheet["C{}".format(i)].value
		Asunto=sheet["D{}".format(i)].value

		if(OC):
			OS_arr.append(OS)
			Cliente_arr.append(Cliente)
			OC_arr.append(OC)
			Asunto_arr.append(Asunto)
		else:
			OS_arr_sinOC.append(OS)
			Cliente_arr_sinOC.append(Cliente)
			OC_arr_sinOC.append(OC)
			Asunto_arr_sinOC.append(Asunto)

def escribir_doc(mes,cierre):

	for i in range(len(OS_arr)):
		document=Document()
		font=document.styles['Normal'].font
		font.name='Calibri(Body)'
		font.size=Pt(11)
		document

		p = document.add_paragraph()
		p.alignment = WD_ALIGN_PARAGRAPH.CENTER	
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.add_run('Memorándum de Completamiento de Servicios \n \n')

		p = document.add_paragraph()
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.add_run('Para: Contabilidad\nDe:    KRCP\nAsunto: {}\n'.format(Asunto_arr[i]))

		p = document.add_paragraph()
		p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.add_run('En relación con la OS {} mediante el presente informamos que a la fecha de este Memorándum el servicio ha sido completado en el periodo: {} según las especificaciones acordadas en la OC {} con el cliente {}. De acuerdo con nuestra experiencia en servicios similares, no existan labores adicionales a las ya realizadas, quedando solo a la espera del recojo por parte del cliente mencionado.'.format(Asunto_arr[i],mes,OC_arr[i],Cliente_arr[i]))


		p = document.add_paragraph()
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.add_run('\n\n______________________________\nFirma\nNombre: \nCargo:     \nFecha:    {}'.format(cierre))

		p = document.add_paragraph()
		run = p.add_run()
		run.add_break(WD_BREAK.PAGE)

		p = document.add_paragraph()
		p.alignment = WD_ALIGN_PARAGRAPH.CENTER	
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.add_run('Memorándum de Completamiento de Servicios \n \n')

		p = document.add_paragraph()
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.add_run('Para: \nDe:    \nAsunto: {}\n'.format(Asunto_arr[i]))

		p = document.add_paragraph()
		p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.add_run('En relación con la OS {} mediante el presente informamos que a la fecha de este Memorándum el servicio ha sido completado en el periodo: {} según las especificaciones acordadas en la OC {} con el cliente {}. De acuerdo con nuestra experiencia en servicios similares, no existan labores adicionales a las ya realizadas, quedando solo a la espera del recojo por parte del cliente mencionado.'.format(Asunto_arr[i],mes,OC_arr[i],Cliente_arr[i]))


		p = document.add_paragraph()
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.add_run('\n\n______________________________\nFirma\nNombre: \nCargo:     \nFecha:   {}'.format(cierre))

		p = document.add_paragraph()
		run = p.add_run()
		run.add_break(WD_BREAK.PAGE)

		p = document.add_paragraph()
		p.alignment = WD_ALIGN_PARAGRAPH.CENTER	
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.add_run('Memorándum de Completamiento de Servicios \n \n')

		p = document.add_paragraph()
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.add_run('Para: \nDe:    \nAsunto: {}\n'.format(Asunto_arr[i]))

		p = document.add_paragraph()
		p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.add_run('En relación con la OS {} mediante el presente informamos que a la fecha de este Memorándum el servicio ha sido completado en el periodo: {} según las especificaciones acordadas en la OC {} con el cliente {}. De acuerdo con nuestra experiencia en servicios similares, no existan labores adicionales a las ya realizadas, quedando solo a la espera del recojo por parte del cliente mencionado.'.format(Asunto_arr[i],mes,OC_arr[i],Cliente_arr[i]))


		p = document.add_paragraph()
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.add_run('\n\n______________________________\nFirma\nNombre: \nCargo:     \nFecha:    {}'.format(cierre))

		core_properties = document.core_properties
		core_properties.author = 'Eric'

		document.save('Memorandum de Completamiento de Servicios- - {}.docx'.format(OS_arr[i]))

	for i in range(len(OS_arr_sinOC)):
		document=Document()
		font=document.styles['Normal'].font
		font.name='Calibri(Body)'
		font.size=Pt(11)
		document

		p = document.add_paragraph()
		p.alignment = WD_ALIGN_PARAGRAPH.CENTER	
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.add_run('Memorándum de Completamiento de Servicios \n \n')

		p = document.add_paragraph()
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.add_run('Para: \nDe:    \nAsunto: {}\n'.format(Asunto_arr_sinOC[i]))

		p = document.add_paragraph()
		p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.add_run('En relación con la OS {} mediante el presente informamos que a la fecha de este Memorándum el servicio ha sido completado en el periodo: {} según los tiempos de reparación comprometidos con el cliente {}. De acuerdo con nuestra experiencia en servicios similares, no existan labores adicionales a las ya realizadas, quedando solo a la espera del recojo por parte del cliente mencionado.'.format(Asunto_arr_sinOC[i],mes,Cliente_arr_sinOC[i]))


		p = document.add_paragraph()
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.add_run('\n\n______________________________\nFirma\nNombre: \nCargo:     \nFecha:    {}'.format(cierre))

		p = document.add_paragraph()
		run = p.add_run()
		run.add_break(WD_BREAK.PAGE)

		p = document.add_paragraph()
		p.alignment = WD_ALIGN_PARAGRAPH.CENTER	
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.add_run('Memorándum de Completamiento de Servicios \n \n')

		p = document.add_paragraph()
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.add_run('Para: \nDe:    \nAsunto: {}\n'.format(Asunto_arr_sinOC[i]))

		p = document.add_paragraph()
		p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.add_run('En relación con la OS {} mediante el presente informamos que a la fecha de este Memorándum el servicio ha sido completado en el periodo: {} según los tiempos de reparación comprometidos con el cliente {}. De acuerdo con nuestra experiencia en servicios similares, no existan labores adicionales a las ya realizadas, quedando solo a la espera del recojo por parte del cliente mencionado.'.format(Asunto_arr_sinOC[i],mes,Cliente_arr_sinOC[i]))


		p = document.add_paragraph()
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.add_run('\n\n______________________________\nFirma\nNombre: \nCargo:     \nFecha:    {}'.format(cierre))

		p = document.add_paragraph()
		run = p.add_run()
		run.add_break(WD_BREAK.PAGE)

		p = document.add_paragraph()
		p.alignment = WD_ALIGN_PARAGRAPH.CENTER	
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.add_run('Memorándum de Completamiento de Servicios \n \n')

		p = document.add_paragraph()
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.add_run('Para: \nDe:    \nAsunto: {}\n'.format(Asunto_arr_sinOC[i]))

		p = document.add_paragraph()
		p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.add_run('En relación con la OS {} mediante el presente informamos que a la fecha de este Memorándum el servicio ha sido completado en el periodo: {} según los tiempos de reparación comprometidos con el cliente {}. De acuerdo con nuestra experiencia en servicios similares, no existan labores adicionales a las ya realizadas, quedando solo a la espera del recojo por parte del cliente mencionado.'.format(Asunto_arr_sinOC[i],mes,Cliente_arr_sinOC[i]))


		p = document.add_paragraph()
		p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
		p.add_run('\n\n______________________________\nFirma\nNombre: \nCargo:     \nFecha:    {}'.format(cierre))

		core_properties = document.core_properties
		core_properties.author = 'Eric'

		document.save('Memorandum de Completamiento de Servicios- - {}.docx'.format(OS_arr_sinOC[i]))

def generacion():
  global ruta, periodo,inicio,fin,fecha_cierre
  inicio=int(inicio)
  fin=int(fin)
  #Ruta del Excel
  xlsx_file = Path(ruta)
  #Obtener OS, Clientes, OC y Asuntos del archivo Excel
  fin=fin+1
  leer_xl(xlsx_file,inicio,fin);
  #Generar Memorandums
  escribir_doc(periodo,fecha_cierre)

def main():

  root = tk.Tk()
  root.title("KRCP-MEMOS")
  root.iconbitmap("krcp.ico")
  frame = tk.Frame(root)
  root.geometry('500x200')
  ents = makeform(root, fields)


  b1 = tk.Button(root,text='Generar', command=(lambda e=root: generar(root,ents)))
  b1.pack(side=tk.LEFT, padx=5, pady=5)
  
  b2 = tk.Button(root, text='Salir', command=(lambda e=root: cerrar(root)))
  b2.pack(side=tk.LEFT, padx=5, pady=5) 
  
  root.mainloop()

if __name__ == '__main__':
	main()